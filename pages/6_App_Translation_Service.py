import streamlit as st
import os
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import gtts
from io import BytesIO
import requests
import json
from google.cloud import translate_v2 as translate
from google.auth.api_key import Credentials

# Set Tesseract path (change this according to your system)
#pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Windows example
# For Linux/Mac: pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# Set up Gemini API
API_KEY = os.getenv("GOOGLE_TRANSLATION_API_KEY")
if not API_KEY:
    st.error('GOOGLE_TRANSLATION_API_KEY environment variable not set.')
    st.stop()


# Function to extract text from image
def image_to_text(image_path):
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    return text

# Function to extract text from PDF
def pdf_to_text(pdf_path):
    images = convert_from_path(pdf_path)
    text = ""
    for img in images:
        text += pytesseract.image_to_string(img)
    return text

def translate_text(text, target_language):
    """Translates text into the target language using an API key."""
    translate_client = translate.Client(credentials=Credentials(API_KEY))

    translation = translate_client.translate(
        text,
        target_language=target_language
    )

    return translation['translatedText']


# Function to create and save DOCX
def create_docx1(text, filename):
    doc = Document()
    doc.add_paragraph(text)
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def create_docx(text, filename):
    # Sanitize text to be XML compatible
    # Remove or replace problematic characters
    sanitized_text = ""
    for char in text:
        # XML can't handle control characters except tab, newline, and carriage return
        if ord(char) >= 32 or char in ['\t', '', '\r']:
            sanitized_text += char
        else:
            # Replace control characters with a space or empty string
            sanitized_text += " "
    
    # Create the document with sanitized text
    doc = Document()
    doc.add_paragraph(sanitized_text)
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Streamlit app
def main():
    st.title("Notice Translation App")
    st.write("Upload a notice (PDF or image) to translate to Hindi or Kannada")

    # File upload
    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'png', 'jpg', 'jpeg'])

    # Language selection
    options = ["Hindi", "Kannada"]
    selected_lang = st.selectbox("Choose your target language:", options, index=None)
    st.write(f"You selected: {selected_lang}")

    if selected_lang == "Hindi":
        target_lang = "hi"
    if selected_lang == "Kannada":
        target_lang = "ka"
   
    if uploaded_file is not None:
        # Save the uploaded file temporarily
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        temp_file_path = f"temp{file_ext}"
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        try:
            # Extract text based on file type
            if file_ext == '.pdf':
                extracted_text = pdf_to_text(temp_file_path)
            else:
                extracted_text = image_to_text(temp_file_path)

            if extracted_text.strip():
                st.subheader("Original Text")
                st.text_area("Original Text", extracted_text, height=200, key="original")

                # Translate text
                with st.spinner('Translating...'):
                    translated_text = translate_text(extracted_text, target_lang)

                st.subheader("Translated Text")
                st.text_area("Translated Text", translated_text, height=200, key="translated")

                # Create download buttons
                col1 = st.columns(1)[0]

                with col1:
                    # Download DOCX
                    docx_file = create_docx(translated_text, "translated_notice.docx")
                    st.download_button(
                        label="Download as DOCX",
                        data=docx_file,
                        file_name="translated_notice.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )             
            else:
                st.warning("No text could be extracted from the file. Please try with a clearer image or PDF.")

        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)

if __name__ == "__main__":
    main()
